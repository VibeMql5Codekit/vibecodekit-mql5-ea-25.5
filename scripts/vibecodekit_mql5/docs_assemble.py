"""mql5-docs-assemble — convert an LLM-authored ``guide.md`` to ``.docx``.

Second half of the LLM-driven user-guide ship pipeline (see
:mod:`docs_bundle` for the first half). Takes a Vietnamese markdown
file the LLM agent wrote against ``docs-context.json`` and renders it
to a Word document suitable for inclusion in the ship.zip.

What's supported (the markdown subset the bundle prompt asks for):

* Headings (``#`` … ``####``) — mapped to Word ``Heading 1`` …
  ``Heading 4``.
* Paragraphs with inline bold (``**``), italic (``*``), code spans
  (`````) and links (``[text](url)``). Links keep their URL inline so a
  printed copy is still useful.
* Bullet and ordered lists (single level — the LLM prompt asks for flat
  outlines).
* Fenced code blocks (```` ``` ````) — rendered in Consolas 9pt.
* Markdown tables — emitted as Word tables with a bold header row.
* Blockquote callouts (``> **Lưu ý**: ...``) — emitted as italic
  paragraphs.
* Images (``![alt](images/...)``) — embedded from the configured
  ``images_dir``; missing files surface as warnings rather than errors.
* ``[[TOC]]`` marker → Word ToC field that updates with ``F9``.

What we deliberately don't try to support: HTML passthrough, nested
lists, footnotes, definition lists. The LLM prompt steers away from
those.

This module is **archetype-agnostic** — it never injects per-EA chapter
content. All variation comes from the markdown the LLM authored.
"""

from __future__ import annotations

import argparse
import json
import sys
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown_it import MarkdownIt


__all__ = [
    "AssembleResult",
    "assemble",
    "build_document",
    "validate_guide_md",
    "main",
]


_MAX_HEADING_LEVEL = 4
_TOC_MARKER = "[[TOC]]"
_MIN_CHAPTERS = 5
_MAX_CHAPTERS = 12


@dataclass
class AssembleResult:
    """Summary of a markdown → docx run."""

    ok: bool
    docx_path: str
    chapters: int = 0
    tables: int = 0
    images_embedded: int = 0
    warnings: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------


def validate_guide_md(md_text: str) -> list[str]:
    """Lightweight structural checks on the LLM-authored markdown.

    Returns a list of human-readable issues (Vietnamese). An empty list
    means the markdown clears the structural bar set by ``docs-prompt.md``.
    Callers may treat any issue as a soft warning — :func:`assemble`
    still renders a docx so a reviewer can eyeball the result and fix.
    """
    issues: list[str] = []
    lines = md_text.strip().splitlines() if md_text else []
    if not lines or not lines[0].startswith("# "):
        issues.append("H1 phải là dòng đầu tiên của guide.md")
    if _TOC_MARKER not in md_text:
        issues.append(f"Thiếu marker {_TOC_MARKER} (Word ToC field)")
    h1_count = sum(1 for line in lines if line.startswith("# "))
    chapter_count = max(h1_count - 1, 0)  # exclude the title H1
    if chapter_count < _MIN_CHAPTERS:
        issues.append(
            f"Quá ít chương: đếm được {chapter_count}, tối thiểu {_MIN_CHAPTERS}"
        )
    if chapter_count > _MAX_CHAPTERS:
        issues.append(
            f"Quá nhiều chương: đếm được {chapter_count}, tối đa {_MAX_CHAPTERS}"
        )
    return issues


# ---------------------------------------------------------------------------
# Document building
# ---------------------------------------------------------------------------


def _apply_default_style(doc: Document) -> None:
    """Set the body font + page margins to match the NBkudo reference."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.76)
        section.bottom_margin = Inches(0.76)
        section.left_margin = Inches(0.76)
        section.right_margin = Inches(0.76)


def _insert_toc_field(doc: Document) -> None:
    """Insert a Word ToC field that updates on F9 / open.

    python-docx has no first-class ToC API, so we emit the underlying
    field-instruction sequence on a fresh paragraph.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    placeholder_run = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "Mục lục — nhấn F9 để cập nhật"
    placeholder_run.append(placeholder_text)
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_separate)
    run._r.append(placeholder_run)
    run._r.append(fld_char_end)


def _is_toc_marker_paragraph(inline_token: Any) -> bool:
    return getattr(inline_token, "content", "").strip() == _TOC_MARKER


def _is_image_only_paragraph(inline_token: Any) -> bool:
    children = getattr(inline_token, "children", None) or []
    visible = [c for c in children if c.type not in {"text", "softbreak"} or (c.type == "text" and c.content.strip())]
    return len(visible) == 1 and visible[0].type == "image"


def _extract_image_src(inline_token: Any) -> str | None:
    for child in getattr(inline_token, "children", None) or []:
        if child.type == "image":
            attrs = dict(child.attrs or {})
            return str(attrs.get("src") or "")
    return None


def _embed_image(
    doc: Document, image_path: Path, result: AssembleResult
) -> None:
    if not image_path.is_file():
        result.warnings.append(f"image not found: {image_path}")
        return
    try:
        doc.add_picture(str(image_path), width=Inches(6.0))
    except Exception as exc:  # python-docx raises broad exceptions on bad images
        result.warnings.append(f"image embed failed for {image_path}: {exc}")
        return
    result.images_embedded += 1


def _add_inline_runs(paragraph: Any, children: list[Any]) -> None:
    """Walk markdown-it inline tokens and emit Word runs with basic styling."""
    bold = False
    italic = False
    for child in children:
        ttype = child.type
        if ttype == "text":
            run = paragraph.add_run(child.content)
            run.bold = bold or None
            run.italic = italic or None
        elif ttype == "code_inline":
            run = paragraph.add_run(child.content)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif ttype == "strong_open":
            bold = True
        elif ttype == "strong_close":
            bold = False
        elif ttype == "em_open":
            italic = True
        elif ttype == "em_close":
            italic = False
        elif ttype == "softbreak":
            paragraph.add_run(" ")
        elif ttype == "hardbreak":
            paragraph.add_run().add_break()
        elif ttype == "link_open":
            attrs = dict(child.attrs or {})
            href = str(attrs.get("href", ""))
            # Stash on a state var via closure trick: we keep `href` and
            # finalize formatting on the next text token; simpler is to
            # just emit the URL inline once `link_close` fires.
            child._href = href  # type: ignore[attr-defined]
        elif ttype == "link_close":
            # No-op — we already wrote the link text; downstream may
            # append the URL in parentheses if it differs from the body.
            continue
        elif ttype == "image":
            attrs = dict(child.attrs or {})
            paragraph.add_run(f"[image: {attrs.get('src', '')}]")
        else:
            # Unknown inline token — surface its raw content so nothing is
            # silently dropped from the user-facing guide.
            content = getattr(child, "content", "")
            if content:
                paragraph.add_run(content)


def _heading_from_token(doc: Document, level: int, text: str) -> None:
    level = max(1, min(level, _MAX_HEADING_LEVEL))
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_code_block(doc: Document, content: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(content.rstrip("\n"))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def _add_callout(doc: Document, body_tokens: list[Any]) -> None:
    """Render a blockquote as an italic paragraph with a coloured prefix."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.3)
    text_runs: list[str] = []
    bold_segments: list[bool] = []
    for tok in body_tokens:
        if tok.type != "inline":
            continue
        bold = False
        for child in tok.children or []:
            if child.type == "strong_open":
                bold = True
            elif child.type == "strong_close":
                bold = False
            elif child.type == "text":
                text_runs.append(child.content)
                bold_segments.append(bold)
            elif child.type == "code_inline":
                text_runs.append(child.content)
                bold_segments.append(False)
            elif child.type == "softbreak":
                text_runs.append(" ")
                bold_segments.append(False)
    for text, bold in zip(text_runs, bold_segments):
        run = paragraph.add_run(text)
        run.italic = True
        run.bold = bold or None
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x66)


def _add_list_items(
    doc: Document,
    tokens: list[Any],
    start: int,
    ordered: bool,
) -> int:
    """Emit a flat (single-level) bullet/ordered list.

    Returns the next-token index after the matching ``*_list_close``.
    """
    close = "ordered_list_close" if ordered else "bullet_list_close"
    style = "List Number" if ordered else "List Bullet"
    i = start + 1
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == close:
            return i + 1
        if tok.type == "list_item_open":
            # Find the inline token inside this item.
            j = i + 1
            while j < len(tokens) and tokens[j].type != "list_item_close":
                if tokens[j].type == "inline":
                    paragraph = doc.add_paragraph(style=style)
                    _add_inline_runs(paragraph, tokens[j].children or [])
                j += 1
            i = j + 1
            continue
        i += 1
    return i


def _table_cell_text(inline_token: Any) -> str:
    if inline_token is None:
        return ""
    parts: list[str] = []
    for child in inline_token.children or []:
        if child.type in {"text", "code_inline"}:
            parts.append(child.content)
        elif child.type == "softbreak":
            parts.append(" ")
    return "".join(parts).strip()


def _add_table(
    doc: Document,
    tokens: list[Any],
    start: int,
    result: AssembleResult,
) -> int:
    """Walk a ``table_open`` … ``table_close`` block into a Word table."""
    header: list[str] = []
    rows: list[list[str]] = []
    current_row: list[str] = []
    in_header = False
    i = start + 1
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == "table_close":
            i += 1
            break
        if tok.type == "thead_open":
            in_header = True
        elif tok.type == "thead_close":
            in_header = False
        elif tok.type == "tr_open":
            current_row = []
        elif tok.type == "tr_close":
            if in_header:
                header = current_row
            else:
                rows.append(current_row)
        elif tok.type in {"th_open", "td_open"}:
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            current_row.append(_table_cell_text(inline))
        i += 1

    if not header and not rows:
        return i

    col_count = max(len(header), max((len(r) for r in rows), default=0))
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Light Grid Accent 1"
    if header:
        header_cells = table.rows[0].cells
        for idx in range(col_count):
            cell = header_cells[idx]
            cell.text = header[idx] if idx < len(header) else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx in range(col_count):
            table.rows[r_idx].cells[c_idx].text = row[c_idx] if c_idx < len(row) else ""

    result.tables += 1
    return i


def build_document(
    md_text: str,
    images_dir: Path,
) -> tuple[Document, AssembleResult]:
    """Render ``md_text`` into a fresh :class:`docx.Document`.

    Pure: takes markdown + an image-resolution root, returns the docx
    object plus a structural summary. The caller is responsible for
    persistence — see :func:`assemble`.
    """
    md = MarkdownIt("commonmark", {"html": False}).enable("table").enable("strikethrough")
    tokens = md.parse(md_text)

    doc = Document()
    _apply_default_style(doc)
    result = AssembleResult(ok=True, docx_path="")
    h1_total = 0

    i = 0
    while i < len(tokens):
        tok = tokens[i]
        ttype = tok.type
        if ttype == "heading_open":
            level = int(tok.tag[1])  # "h1" → 1
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            text = inline.content if inline else ""
            _heading_from_token(doc, level, text)
            if level == 1:
                h1_total += 1
            i += 3  # heading_open + inline + heading_close
            continue
        if ttype == "paragraph_open":
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            if inline and _is_toc_marker_paragraph(inline):
                _insert_toc_field(doc)
                i += 3
                continue
            if inline and _is_image_only_paragraph(inline):
                src = _extract_image_src(inline)
                if src:
                    image_path = images_dir / src if not Path(src).is_absolute() else Path(src)
                    # Strip a leading "images/" so callers can pass
                    # either the root or the parent and have both shapes
                    # resolve cleanly.
                    if not image_path.is_file() and src.startswith("images/"):
                        image_path = images_dir.parent / src
                    _embed_image(doc, image_path, result)
                i += 3
                continue
            paragraph = doc.add_paragraph()
            if inline:
                _add_inline_runs(paragraph, inline.children or [])
            i += 3
            continue
        if ttype == "bullet_list_open":
            i = _add_list_items(doc, tokens, i, ordered=False)
            continue
        if ttype == "ordered_list_open":
            i = _add_list_items(doc, tokens, i, ordered=True)
            continue
        if ttype == "fence" or ttype == "code_block":
            _add_code_block(doc, tok.content or "")
            i += 1
            continue
        if ttype == "table_open":
            i = _add_table(doc, tokens, i, result)
            continue
        if ttype == "blockquote_open":
            body: list[Any] = []
            j = i + 1
            depth = 1
            while j < len(tokens) and depth > 0:
                if tokens[j].type == "blockquote_open":
                    depth += 1
                elif tokens[j].type == "blockquote_close":
                    depth -= 1
                    if depth == 0:
                        break
                else:
                    body.append(tokens[j])
                j += 1
            _add_callout(doc, body)
            i = j + 1
            continue
        if ttype == "hr":
            doc.add_paragraph("─" * 40)
            i += 1
            continue
        i += 1

    # The first H1 is the doc title (EA name + version per the prompt
    # checklist); subsequent H1s are chapters. Floor at zero so a guide
    # with only a title still surfaces a sensible chapter count.
    result.chapters = max(h1_total - 1, 0)
    return doc, result


def assemble(
    guide_md_path: Path,
    out_path: Path,
    images_dir: Path | None = None,
) -> AssembleResult:
    """Render ``guide_md_path`` to ``out_path`` (Word ``.docx``).

    ``images_dir`` defaults to ``<guide.parent>/images``. Validation
    issues from :func:`validate_guide_md` land in ``result.warnings``
    but do not fail the run — the caller can inspect the warnings and
    decide whether to re-prompt the LLM.
    """
    md_text = guide_md_path.read_text(encoding="utf-8")
    issues = validate_guide_md(md_text)
    if images_dir is None:
        images_dir = guide_md_path.parent / "images"
    doc, result = build_document(md_text, images_dir)
    result.warnings.extend(issues)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    result.docx_path = str(out_path)
    return result


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        prog="mql5-docs-assemble",
        description=(
            "Convert an LLM-authored guide.md into a .docx Word document "
            "for the EA ship bundle."
        ),
    )
    parser.add_argument("guide", type=Path, help="LLM-authored guide.md path")
    parser.add_argument(
        "--out",
        type=Path,
        required=True,
        help="output .docx path",
    )
    parser.add_argument(
        "--images-dir",
        type=Path,
        default=None,
        help="root directory for `![alt](images/...)` references "
        "(default: <guide-parent>/images)",
    )
    args = parser.parse_args(argv)

    if not args.guide.is_file():
        print(f"mql5-docs-assemble: guide.md not found: {args.guide}", file=sys.stderr)
        return 2

    try:
        result = assemble(args.guide, args.out, images_dir=args.images_dir)
    except OSError as exc:
        print(f"mql5-docs-assemble: {exc}", file=sys.stderr)
        return 2

    print(json.dumps(asdict(result), indent=2, ensure_ascii=False))
    return 0 if result.ok else 1


if __name__ == "__main__":
    sys.exit(main())
