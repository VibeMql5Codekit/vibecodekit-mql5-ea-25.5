"""Markdown-token → python-docx rendering helpers for ``docs_assemble``.

Split out from :mod:`docs_assemble` to keep that module under the
audit script's per-module LOC ceiling (see ``MODULE_LOC_CEILING`` in
``scripts/audit-plan-v5.py``). All helpers in this module are
**private**: the only public entrypoint to the markdown-to-docx
pipeline remains :func:`docs_assemble.assemble`.

The helpers are deliberately small and stateless (they take a
``Document`` plus a token list) so the orchestration in
:func:`docs_assemble.build_document` stays linear and easy to audit.
"""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING, Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

if TYPE_CHECKING:
    from .docs_assemble import AssembleResult


MAX_HEADING_LEVEL = 4
TOC_MARKER = "[[TOC]]"


def apply_default_style(doc: Document) -> None:
    """Set the body font + page margins to match the NBkudo reference."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.76)
        section.bottom_margin = Inches(0.76)
        section.left_margin = Inches(0.76)
        section.right_margin = Inches(0.76)


def insert_toc_field(doc: Document) -> None:
    """Insert a Word ToC field that updates on F9 / open.

    python-docx has no first-class ToC API, so we emit the underlying
    field-instruction sequence on a fresh paragraph.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    placeholder_run = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "Mục lục — nhấn F9 để cập nhật"
    placeholder_run.append(placeholder_text)
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_separate)
    run._r.append(placeholder_run)
    run._r.append(fld_char_end)


def is_toc_marker_paragraph(inline_token: Any) -> bool:
    return getattr(inline_token, "content", "").strip() == TOC_MARKER


def is_image_only_paragraph(inline_token: Any) -> bool:
    children = getattr(inline_token, "children", None) or []
    visible = [c for c in children if c.type not in {"text", "softbreak"} or (c.type == "text" and c.content.strip())]
    return len(visible) == 1 and visible[0].type == "image"


def extract_image_src(inline_token: Any) -> str | None:
    for child in getattr(inline_token, "children", None) or []:
        if child.type == "image":
            attrs = dict(child.attrs or {})
            return str(attrs.get("src") or "")
    return None


def embed_image(doc: Document, image_path: Path, result: "AssembleResult") -> None:
    if not image_path.is_file():
        result.warnings.append(f"image not found: {image_path}")
        return
    try:
        doc.add_picture(str(image_path), width=Inches(6.0))
    except Exception as exc:  # python-docx raises broad exceptions on bad images
        result.warnings.append(f"image embed failed for {image_path}: {exc}")
        return
    result.images_embedded += 1


def add_inline_runs(paragraph: Any, children: list[Any]) -> None:
    """Walk markdown-it inline tokens and emit Word runs with basic styling.

    Links keep their URL inline: ``[text](url)`` renders as ``text
    (url)`` so a printed copy of the document remains useful when the
    hyperlink can't be clicked.
    """
    bold = False
    italic = False
    pending_href = ""
    for child in children:
        ttype = child.type
        if ttype == "text":
            run = paragraph.add_run(child.content)
            run.bold = bold or None
            run.italic = italic or None
        elif ttype == "code_inline":
            run = paragraph.add_run(child.content)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif ttype == "strong_open":
            bold = True
        elif ttype == "strong_close":
            bold = False
        elif ttype == "em_open":
            italic = True
        elif ttype == "em_close":
            italic = False
        elif ttype == "softbreak":
            paragraph.add_run(" ")
        elif ttype == "hardbreak":
            paragraph.add_run().add_break()
        elif ttype == "link_open":
            pending_href = str(dict(child.attrs or {}).get("href", ""))
        elif ttype == "link_close":
            if pending_href:
                url_run = paragraph.add_run(f" ({pending_href})")
                url_run.font.size = Pt(9)
                url_run.italic = True
            pending_href = ""
        elif ttype == "image":
            attrs = dict(child.attrs or {})
            paragraph.add_run(f"[image: {attrs.get('src', '')}]")
        else:
            content = getattr(child, "content", "")
            if content:
                paragraph.add_run(content)


def heading_from_token(doc: Document, level: int, text: str) -> None:
    level = max(1, min(level, MAX_HEADING_LEVEL))
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_code_block(doc: Document, content: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(content.rstrip("\n"))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_callout(doc: Document, body_tokens: list[Any]) -> None:
    """Render a blockquote as an italic paragraph with a coloured prefix."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.3)
    text_runs: list[str] = []
    bold_segments: list[bool] = []
    for tok in body_tokens:
        if tok.type != "inline":
            continue
        bold = False
        for child in tok.children or []:
            if child.type == "strong_open":
                bold = True
            elif child.type == "strong_close":
                bold = False
            elif child.type == "text":
                text_runs.append(child.content)
                bold_segments.append(bold)
            elif child.type == "code_inline":
                text_runs.append(child.content)
                bold_segments.append(False)
            elif child.type == "softbreak":
                text_runs.append(" ")
                bold_segments.append(False)
    for text, bold in zip(text_runs, bold_segments, strict=False):
        run = paragraph.add_run(text)
        run.italic = True
        run.bold = bold or None
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x66)


def add_list_items(
    doc: Document,
    tokens: list[Any],
    start: int,
    ordered: bool,
) -> int:
    """Emit a flat (single-level) bullet/ordered list.

    Returns the next-token index after the matching ``*_list_close``.
    """
    close = "ordered_list_close" if ordered else "bullet_list_close"
    style = "List Number" if ordered else "List Bullet"
    i = start + 1
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == close:
            return i + 1
        if tok.type == "list_item_open":
            j = i + 1
            while j < len(tokens) and tokens[j].type != "list_item_close":
                if tokens[j].type == "inline":
                    paragraph = doc.add_paragraph(style=style)
                    add_inline_runs(paragraph, tokens[j].children or [])
                j += 1
            i = j + 1
            continue
        i += 1
    return i


def _table_cell_text(inline_token: Any) -> str:
    if inline_token is None:
        return ""
    parts: list[str] = []
    for child in inline_token.children or []:
        if child.type in {"text", "code_inline"}:
            parts.append(child.content)
        elif child.type == "softbreak":
            parts.append(" ")
    return "".join(parts).strip()


def add_table(
    doc: Document,
    tokens: list[Any],
    start: int,
    result: "AssembleResult",
) -> int:
    """Walk a ``table_open`` … ``table_close`` block into a Word table."""
    header: list[str] = []
    rows: list[list[str]] = []
    current_row: list[str] = []
    in_header = False
    i = start + 1
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == "table_close":
            i += 1
            break
        if tok.type == "thead_open":
            in_header = True
        elif tok.type == "thead_close":
            in_header = False
        elif tok.type == "tr_open":
            current_row = []
        elif tok.type == "tr_close":
            if in_header:
                header = current_row
            else:
                rows.append(current_row)
        elif tok.type in {"th_open", "td_open"}:
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            current_row.append(_table_cell_text(inline))
        i += 1

    if not header and not rows:
        return i

    col_count = max(len(header), max((len(r) for r in rows), default=0))
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Light Grid Accent 1"
    if header:
        header_cells = table.rows[0].cells
        for idx in range(col_count):
            cell = header_cells[idx]
            cell.text = header[idx] if idx < len(header) else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx in range(col_count):
            table.rows[r_idx].cells[c_idx].text = row[c_idx] if c_idx < len(row) else ""

    result.tables += 1
    return i
